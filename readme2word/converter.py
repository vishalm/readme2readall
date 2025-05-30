import base64
import os
import re
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import markdown
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image


class ReadmeToWordConverter:
    def __init__(self):
        self.stats = {
            "headings": 0,
            "tables": 0,
            "code_blocks": 0,
            "mermaid_diagrams": 0,
            "images": 0,
        }
        self.mermaid_counter = 0
        self.debug_mode = True  # Enable debug output

    def convert(
        self,
        readme_content: str,
        output_filename: str,
        include_toc: bool = True,
        diagram_style: str = "default",
    ) -> str:
        """Convert README content to Word document"""
        # Reset stats
        self.stats = {k: 0 for k in self.stats.keys()}
        self.mermaid_counter = 0

        if self.debug_mode:
            print(f"🔍 Starting conversion with diagram style: {diagram_style}")
            print(f"📝 Content length: {len(readme_content)} characters")

        # Create new document
        doc = Document()

        # Set up document styles
        self._setup_document_styles(doc)

        # Add title
        title = self._extract_title(readme_content)
        if title:
            doc.add_heading(title, 0)

        # Add table of contents placeholder if requested
        if include_toc:
            self._add_table_of_contents(doc)

        # Process mermaid diagrams first (convert to images)
        if self.debug_mode:
            mermaid_count = len(
                re.findall(r"```mermaid\n(.*?)\n```", readme_content, re.DOTALL)
            )
            print(f"🎨 Found {mermaid_count} Mermaid diagrams to convert")

        content_with_images = self._process_mermaid_diagrams(
            readme_content, diagram_style
        )

        # Convert markdown to HTML
        md = markdown.Markdown(extensions=["tables", "fenced_code", "codehilite"])
        html_content = md.convert(content_with_images)

        # Parse HTML and convert to Word
        soup = BeautifulSoup(html_content, "html.parser")
        self._convert_html_to_word(soup, doc)

        # Save document
        # Check if filename already has .docx extension to avoid double extension
        if output_filename.endswith(".docx"):
            output_path = Path(output_filename)
        else:
            output_path = Path(f"{output_filename}.docx")

        # Create parent directories if they don't exist
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc.save(str(output_path))

        if self.debug_mode:
            print(f"✅ Document saved to: {output_path}")
            print(f"📊 Final stats: {self.stats}")

        return str(output_path)

    def _setup_document_styles(self, doc: Document) -> None:
        """Set up custom styles for the document"""
        styles = doc.styles

        # Code block style
        try:
            code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = "Consolas"
            code_style.font.size = Pt(9)
            code_style.paragraph_format.left_indent = Inches(0.5)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
        except BaseException:
            pass  # Style might already exist

    def _extract_title(self, content: str) -> str:
        """Extract the main title from README content"""
        lines = content.split("\n")
        for line in lines:
            if line.strip().startswith("# "):
                return line.strip()[2:].strip()
        return ""

    def _add_table_of_contents(self, doc: Document) -> None:
        """Add a table of contents placeholder"""
        doc.add_heading("Table of Contents", 1)
        p = doc.add_paragraph()
        p.add_run(
            "(Table of contents will be generated when you open the document in Word)"
        )
        p.italic = True
        doc.add_page_break()

    def _process_mermaid_diagrams(self, content: str, style: str = "default") -> str:
        """Convert Mermaid diagrams to images using mermaid.ink API"""
        # Improved regex pattern to handle various whitespace scenarios
        mermaid_pattern = r"```mermaid\s*\n(.*?)\n\s*```"

        def replace_mermaid(match):
            mermaid_code = match.group(1).strip()
            if self.debug_mode:
                print(f"🎨 Processing Mermaid diagram {self.mermaid_counter + 1}:")
                print(f"   Code preview: {mermaid_code[:50]}...")

            try:
                # Create image from mermaid code
                image_path = self._mermaid_to_image(mermaid_code, style)
                if image_path:
                    self.stats["mermaid_diagrams"] += 1
                    if self.debug_mode:
                        print(f"   ✅ Successfully converted to: {image_path}")
                    # Return markdown image syntax with absolute path
                    return (
                        f"\n![Mermaid Diagram {self.mermaid_counter}]({image_path})\n"
                    )
                else:
                    if self.debug_mode:
                        print(f"   ❌ Failed to convert, falling back to code block")
                    # Fallback to code block if conversion fails
                    return f"\n```\n{mermaid_code}\n```\n"
            except Exception as e:
                if self.debug_mode:
                    print(f"   ❌ Exception during conversion: {e}")
                return f"\n```\n{mermaid_code}\n```\n"

        result = re.sub(mermaid_pattern, replace_mermaid, content, flags=re.DOTALL)

        if self.debug_mode:
            print(
                f"🔄 Mermaid processing complete. Converted {self.stats['mermaid_diagrams']} diagrams"
            )

        return result

    def _mermaid_to_image(
        self, mermaid_code: str, style: str = "default"
    ) -> Optional[str]:
        """Convert mermaid code to image using mermaid.ink API"""
        try:
            self.mermaid_counter += 1

            if self.debug_mode:
                print(f"   🌐 Calling Mermaid API for diagram {self.mermaid_counter}")

            # Clean and encode mermaid code
            cleaned_code = mermaid_code.strip()
            encoded = base64.urlsafe_b64encode(cleaned_code.encode("utf-8")).decode(
                "ascii"
            )

            # Get image from mermaid.ink with proper theme mapping
            theme_map = {
                "default": "default",
                "neutral": "neutral",
                "dark": "dark",
                "forest": "forest",
            }
            actual_theme = theme_map.get(style, "default")

            url = f"https://mermaid.ink/img/{encoded}?theme={actual_theme}"

            if self.debug_mode:
                print(f"   📡 API URL: {url[:80]}...")

            # Make request with better error handling
            headers = {"User-Agent": "README-to-Word-Converter/1.0"}
            response = requests.get(url, timeout=15, headers=headers)

            if self.debug_mode:
                print(
                    f"   📊 Response: {response.status_code}, Content-Type: {response.headers.get('content-type', 'unknown')}"
                )

            if response.status_code == 200:
                # Verify it's actually an image
                content_type = response.headers.get("content-type", "")
                if "image" not in content_type:
                    if self.debug_mode:
                        print(f"   ⚠️  Warning: Expected image but got {content_type}")
                        print(f"   Response preview: {response.text[:100]}...")
                    return None

                # Save image with absolute path
                output_dir = Path("images")
                output_dir.mkdir(parents=True, exist_ok=True)
                image_path = output_dir / f"mermaid_{self.mermaid_counter}.png"

                # Convert to PNG if needed
                try:
                    # Save the response content
                    with open(image_path, "wb") as f:
                        f.write(response.content)

                    # Verify the image can be opened
                    with Image.open(image_path) as img:
                        # Convert to PNG if it's not already
                        if img.format != "PNG":
                            png_path = image_path.with_suffix(".png")
                            img.save(png_path, "PNG")
                            if png_path != image_path:
                                image_path.unlink()  # Remove original
                                image_path = png_path

                    if self.debug_mode:
                        print(
                            f"   💾 Image saved: {image_path} ({image_path.stat().st_size} bytes)"
                        )

                    # Return absolute path for better compatibility
                    return str(image_path.resolve())

                except Exception as img_error:
                    if self.debug_mode:
                        print(f"   ❌ Image processing error: {img_error}")
                    return None
            else:
                if self.debug_mode:
                    print(
                        f"   ❌ API Error {response.status_code}: {response.text[:200]}"
                    )
                return None

        except requests.exceptions.Timeout:
            if self.debug_mode:
                print(f"   ⏰ Timeout error - API took too long to respond")
            return None
        except requests.exceptions.ConnectionError:
            if self.debug_mode:
                print(f"   🌐 Connection error - Check internet connection")
            return None
        except Exception as e:
            if self.debug_mode:
                print(f"   ❌ Unexpected error: {e}")
            return None

    def _convert_html_to_word(self, soup: BeautifulSoup, doc: Document) -> None:
        """Convert HTML elements to Word document elements"""
        for element in soup.children:
            if hasattr(element, "name"):
                self._process_element(element, doc)

    def _process_element(
        self, element: Any, doc: Document, parent_paragraph: Any = None
    ) -> None:
        """Process individual HTML elements"""
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            # Skip h1 if it's the first heading (already used as title)
            if not (level == 1 and self.stats["headings"] == 0):
                doc.add_heading(element.get_text().strip(), level)
            self.stats["headings"] += 1

        elif element.name == "p":
            # Process paragraph content in order, handling mixed text and
            # images
            self._process_paragraph_with_mixed_content(element, doc)

        elif element.name == "table":
            self._convert_table(element, doc)

        elif element.name in ["pre", "code"]:
            self._convert_code_block(element, doc)

        elif element.name == "img":
            self._convert_image(element, doc)

        elif element.name in ["ul", "ol"]:
            self._convert_list(element, doc)

        elif element.name == "blockquote":
            p = doc.add_paragraph(element.get_text().strip())
            p.style = "Quote"

        elif element.name in ["div", "span"]:
            # Process children of container elements
            for child in element.children:
                if hasattr(child, "name"):
                    self._process_element(child, doc)

    def _process_paragraph_with_mixed_content(
        self, element: Any, doc: Document
    ) -> None:
        """Process paragraph content that may contain both text and images in sequence"""
        # Check if paragraph contains images
        img_tags = element.find_all("img")

        if not img_tags:
            # No images, process as regular paragraph
            p = doc.add_paragraph()
            self._process_inline_elements(element, p)
            return

        # Paragraph contains images - need to process content in order
        # Split content around images to maintain proper positioning

        # Get all direct children to process in order
        children = list(element.children)
        current_paragraph = None

        for child in children:
            if isinstance(child, str):
                # Text content
                text = child.strip()
                if text:
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph()
                    current_paragraph.add_run(text)
            elif hasattr(child, "name"):
                if child.name == "img":
                    # Finish current paragraph if it exists
                    current_paragraph = None
                    # Add image at document level (in correct position)
                    self._convert_image(child, doc)
                else:
                    # Other inline elements
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph()
                    self._process_single_inline_element(child, current_paragraph)

    def _process_inline_elements(self, element: Any, paragraph: Any) -> None:
        """Process inline elements within a paragraph"""
        for content in element.children:
            if isinstance(content, str):
                paragraph.add_run(content)
            elif hasattr(content, "name"):
                self._process_single_inline_element(content, paragraph)

    def _process_single_inline_element(self, content: Any, paragraph: Any) -> None:
        """Process a single inline element"""
        if content.name == "strong" or content.name == "b":
            run = paragraph.add_run(content.get_text())
            run.bold = True
        elif content.name == "em" or content.name == "i":
            run = paragraph.add_run(content.get_text())
            run.italic = True
        elif content.name == "code":
            run = paragraph.add_run(content.get_text())
            run.font.name = "Consolas"
        elif content.name == "a":
            # Add hyperlink text (basic implementation)
            run = paragraph.add_run(content.get_text())
            run.font.color.rgb = (
                None  # Blue color would need more complex implementation
            )
        else:
            paragraph.add_run(content.get_text())

    def _convert_table(self, table_element: Any, doc: Document) -> None:
        """Convert HTML table to Word table"""
        rows = table_element.find_all("tr")
        if not rows:
            return

        # Count columns
        max_cols = max(len(row.find_all(["td", "th"])) for row in rows)

        # Create Word table
        word_table = doc.add_table(rows=len(rows), cols=max_cols)
        word_table.style = "Table Grid"

        for i, row in enumerate(rows):
            cells = row.find_all(["td", "th"])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    word_table.cell(i, j).text = cell.get_text().strip()
                    # Make header row bold
                    if cell.name == "th":
                        for paragraph in word_table.cell(i, j).paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        self.stats["tables"] += 1

    def _convert_code_block(self, code_element: Any, doc: Document) -> None:
        """Convert code block to Word"""
        code_text = code_element.get_text()
        p = doc.add_paragraph(code_text)
        try:
            p.style = "Code Block"
        except BaseException:
            # Fallback if custom style not available
            p.style = "No Spacing"
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)

        self.stats["code_blocks"] += 1

    def _convert_image(self, img_element: Any, doc: Document) -> None:
        """Convert image to Word"""
        try:
            src = img_element.get("src", "")
            if self.debug_mode:
                print(f"🖼️  Processing image: {src}")

            if src:
                # Handle both relative and absolute paths
                image_path = Path(src)

                # If it's a relative path, make it absolute
                if not image_path.is_absolute():
                    image_path = Path.cwd() / image_path

                if image_path.exists():
                    if self.debug_mode:
                        print(
                            f"   ✅ Image found: {image_path} ({image_path.stat().st_size} bytes)"
                        )

                    # Add image to document with reasonable size
                    try:
                        # Try to add with 6 inch width, but handle oversized
                        # images
                        doc.add_picture(str(image_path), width=Inches(6))
                        self.stats["images"] += 1

                        # Add caption if alt text exists
                        alt_text = img_element.get("alt", "")
                        if alt_text:
                            p = doc.add_paragraph(alt_text)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.italic = True

                        if self.debug_mode:
                            print(f"   ✅ Image successfully added to document")

                    except Exception as img_add_error:
                        if self.debug_mode:
                            print(
                                f"   ❌ Error adding image to document: {img_add_error}"
                            )
                        # Try with smaller size
                        try:
                            doc.add_picture(str(image_path), width=Inches(4))
                            self.stats["images"] += 1
                            if self.debug_mode:
                                print(f"   ✅ Image added with smaller size")
                        except BaseException:
                            if self.debug_mode:
                                print(
                                    f"   ❌ Failed to add image even with smaller size"
                                )
                else:
                    if self.debug_mode:
                        print(f"   ❌ Image file not found: {image_path}")
            else:
                if self.debug_mode:
                    print(f"   ⚠️  No image source provided")

        except Exception as e:
            if self.debug_mode:
                print(f"   ❌ Error processing image: {e}")

    def _convert_list(self, list_element: Any, doc: Document) -> None:
        """Convert HTML list to Word list"""
        items = list_element.find_all("li")
        for item in items:
            p = doc.add_paragraph(item.get_text().strip())
            if list_element.name == "ul":
                p.style = "List Bullet"
            else:  # ol
                p.style = "List Number"

    def get_conversion_stats(self) -> Dict[str, int]:
        """Get statistics about the conversion"""
        return self.stats.copy()

    def set_debug_mode(self, enabled: bool) -> None:
        """Enable or disable debug output"""
        self.debug_mode = enabled

    def test_mermaid_conversion(self, test_code: Optional[str] = None) -> bool:
        """Test Mermaid conversion functionality"""
        if test_code is None:
            test_code = """graph TD
    A[Start] --> B{Decision}
    B -->|Yes| C[Success]
    B -->|No| D[Retry]
    C --> E[End]
    D --> A"""

        print("🧪 Testing Mermaid conversion...")
        print(f"📝 Test code: {test_code[:50]}...")

        try:
            image_path = self._mermaid_to_image(test_code, "default")
            if image_path and Path(image_path).exists():
                print(f"✅ Test successful! Image saved to: {image_path}")
                return True
            else:
                print("❌ Test failed - no image generated")
                return False
        except Exception as e:
            print(f"❌ Test failed with error: {e}")
            return False


def convert_readme_to_word(
    input_file: Union[str, Path],
    output_file: Optional[Union[str, Path]] = None,
    theme: str = "default",
    debug: bool = False,
) -> None:
    """Convert README file to Word document"""
    input_path = Path(input_file)

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_file}")

    # Read the README content
    with open(input_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Create converter instance
    converter = ReadmeToWordConverter()
    converter.set_debug_mode(debug)

    # Determine output filename
    if output_file is None:
        output_filename = input_path.stem
    else:
        output_filename = Path(output_file).stem

    # Convert to Word
    converter.convert(
        readme_content=content,
        output_filename=output_filename,
        include_toc=True,
        diagram_style=theme,
    )
